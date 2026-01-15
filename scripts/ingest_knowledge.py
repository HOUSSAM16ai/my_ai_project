#!/usr/bin/env python3
"""
Knowledge Ingestion Script - The "Superhuman Converter" V2

This script ingests PDF, Word, Excel, Image, or text files, chunks them into semantic units,
and prepares them for the Memory Agent.

Usage:
    python scripts/ingest_knowledge.py <file_path> [--url http://localhost:8000/memories]

Requirements:
    pip install pypdf requests python-docx openpyxl pandas pillow pytesseract pyyaml
"""

import argparse
import json
import os
import sys
import re
from typing import List, Dict, Generator, Iterator

# Lazy imports to avoid startup crash if deps missing
try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    import docx
except ImportError:
    docx = None

try:
    import pandas as pd
except ImportError:
    pd = None

try:
    from PIL import Image
    import pytesseract
except ImportError:
    Image = None
    pytesseract = None

try:
    import yaml
except ImportError:
    yaml = None

# Configuration
CHUNK_SIZE = 1000  # Characters per chunk
OVERLAP = 100      # Overlap between chunks

def read_pdf_generator(path: str) -> Iterator[str]:
    """Yields text from a PDF file page by page to handle large files."""
    if PdfReader is None:
        raise ImportError("pypdf is required for PDF ingestion")

    try:
        reader = PdfReader(path)
        print(f"📖 Reading PDF Stream: {path} ({len(reader.pages)} pages)...")
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                yield text
            if (i + 1) % 10 == 0:
                print(f"  - Processed {i + 1} pages...", end='\r')
        print(f"  - Finished reading {len(reader.pages)} pages.     ")
    except Exception as e:
        print(f"❌ Error reading PDF: {e}")
        sys.exit(1)

def read_docx_generator(path: str) -> Iterator[str]:
    """Yields text from a Word file paragraph by paragraph."""
    if docx is None:
        raise ImportError("python-docx is required for Word ingestion")

    try:
        doc = docx.Document(path)
        print(f"📖 Reading Word Doc Stream: {path} ({len(doc.paragraphs)} paragraphs)...")
        for para in doc.paragraphs:
            if para.text:
                yield para.text + "\n"
    except Exception as e:
        print(f"❌ Error reading DOCX: {e}")
        sys.exit(1)

def read_xlsx_generator(path: str) -> Iterator[str]:
    """Yields text from an Excel file sheet by sheet."""
    if pd is None:
        raise ImportError("pandas and openpyxl are required for Excel ingestion")

    try:
        print(f"📖 Reading Excel Stream: {path}...")
        # Note: pandas read_excel loads full file, but we can iterate sheets
        xl = pd.ExcelFile(path)
        for sheet_name in xl.sheet_names:
            yield f"--- Sheet: {sheet_name} ---\n"
            df = xl.parse(sheet_name)
            yield df.to_string() + "\n\n"
    except Exception as e:
        print(f"❌ Error reading Excel: {e}")
        sys.exit(1)

def read_image_generator(path: str) -> Iterator[str]:
    """Yields text from an Image using OCR."""
    if Image is None or pytesseract is None:
        raise ImportError("pillow and pytesseract are required for Image OCR")

    try:
        print(f"👁️ Reading Image Stream (OCR): {path}...")
        img = Image.open(path)
        text = pytesseract.image_to_string(img)
        if not text.strip():
            print("⚠️ Warning: OCR returned empty text.")
        yield text
    except Exception as e:
        print(f"❌ Error reading Image: {e}")
        sys.exit(1)

def read_text_file_generator(path: str) -> Iterator[str]:
    """Yields text from a plain text file line by line."""
    try:
        print(f"📖 Reading Text Stream: {path}...")
        with open(path, 'r', encoding='utf-8') as f:
            for line in f:
                yield line
    except Exception as e:
        print(f"❌ Error reading file: {e}")
        sys.exit(1)

def read_structured_markdown_generator(path: str) -> Iterator[Dict]:
    """
    Parses a Markdown file with YAML frontmatter/blocks.
    Yields chunks with specific metadata tags.
    Format:
    ---
    metadata_key: value
    ---
    Content...
    """
    if yaml is None:
         # Fallback to plain text if yaml not installed, but we should install it.
         raise ImportError("pyyaml is required for structured markdown")

    try:
        print(f"🧠 Reading Structured Markdown: {path}...")
        with open(path, 'r', encoding='utf-8') as f:
            content = f.read()

        # Split by YAML separators
        # Regex to find --- block ---
        # Note: This loads full file into memory, which is acceptable for text/markdown exams.

        parts = re.split(r'^---\s*$', content, flags=re.MULTILINE)

        current_metadata = {}

        for part in parts:
            part = part.strip()
            if not part:
                continue

            # Check if this part is YAML metadata
            try:
                parsed = yaml.safe_load(part)
                if isinstance(parsed, dict) and 'metadata' in parsed:
                    # It's a metadata block
                    current_metadata = parsed['metadata']
                    continue
                elif isinstance(parsed, dict) and not 'metadata' in parsed:
                    # Maybe just raw k-v pairs?
                    pass
            except yaml.YAMLError:
                pass # Not YAML, treat as text

            # If we are here, it's content
            # Yield content with current metadata
            yield {
                "text": part,
                "metadata": current_metadata.copy()
            }

    except Exception as e:
        print(f"❌ Error reading Structured Markdown: {e}")
        sys.exit(1)


def stream_chunks(text_generator: Iterator[str], chunk_size: int = CHUNK_SIZE, overlap: int = OVERLAP) -> Iterator[str]:
    """Consumes text stream and yields overlapping chunks."""
    buffer = ""
    for text_part in text_generator:
        buffer += text_part
        while len(buffer) >= chunk_size:
            chunk = buffer[:chunk_size]
            yield chunk
            advance = chunk_size - overlap
            buffer = buffer[advance:]
    if buffer:
        yield buffer

def stream_structured_chunks(generator: Iterator[Dict]) -> Iterator[Dict]:
    """
    Takes structured blocks (text + metadata) and chunks the text
    while preserving the metadata for each chunk.
    """
    for block in generator:
        text = block['text']
        metadata = block['metadata']

        # Chunk the text of this block
        # We wrap the text in a list to use the existing stream_chunks logic
        # or just implement simple chunking here to avoid complexity

        start = 0
        text_len = len(text)

        while start < text_len:
            end = min(start + CHUNK_SIZE, text_len)
            chunk_text = text[start:end]

            yield {
                "content": chunk_text,
                "metadata": metadata
            }

            if end == text_len:
                break

            start += (CHUNK_SIZE - OVERLAP)

def main():
    parser = argparse.ArgumentParser(description="Ingest knowledge from files into Memory Agent format (Streaming Mode).")
    parser.add_argument("file", help="Path to the file (PDF, DOCX, XLSX, JPG, PNG, TXT, MD)")
    parser.add_argument("--url", help="Memory Agent API URL (e.g., http://localhost:8001/memories)", default=None)
    parser.add_argument("--dry-run", action="store_true", help="Print payloads without sending", default=False)

    args = parser.parse_args()

    if not os.path.exists(args.file):
        print(f"❌ File not found: {args.file}")
        sys.exit(1)

    ext = os.path.splitext(args.file)[1].lower()
    base_name = os.path.basename(args.file)
    ext_clean = ext.replace('.', '')

    import httpx

    # Special handling for Structured Markdown
    if ext == '.md':
        structured_stream = read_structured_markdown_generator(args.file)
        chunk_stream = stream_structured_chunks(structured_stream)

        print(f"⚡ Processing structured stream...")
        chunk_count = 0
        success_count = 0

        for i, item in enumerate(chunk_stream):
            chunk_count += 1

            # Construct tags from metadata
            metadata = item['metadata']
            tags = ["ingested", f"{ext_clean}_import", base_name]

            # Add metadata as tags "key:value"
            for k, v in metadata.items():
                if isinstance(v, list):
                    for val in v:
                        tags.append(f"{k}:{val}")
                else:
                    tags.append(f"{k}:{v}")

            payload = {
                "content": item['content'],
                "tags": tags
            }

            if args.url and not args.dry_run:
                try:
                    response = httpx.post(args.url, json=payload, timeout=10.0)
                    if response.status_code == 200:
                        success_count += 1
                    else:
                        print(f"  - Chunk {i} failed ❌ ({response.status_code})")
                except Exception as e:
                    print(f"  - Chunk {i} Error: {e}")
            else:
                if i < 3:
                    print(json.dumps(payload, indent=2, ensure_ascii=False))
                elif i == 3:
                    print("... (streaming continues) ...")

        print(f"\n🏁 Completed. Processed {chunk_count} chunks.")
        if args.url and not args.dry_run:
            print(f"🚀 Successfully sent: {success_count}/{chunk_count}")
        return

    # Standard handling for other files
    # 1. Select Generator
    if ext == '.pdf':
        text_stream = read_pdf_generator(args.file)
    elif ext == '.docx':
        text_stream = read_docx_generator(args.file)
    elif ext == '.xlsx':
        text_stream = read_xlsx_generator(args.file)
    elif ext in ['.png', '.jpg', '.jpeg']:
        text_stream = read_image_generator(args.file)
    else:
        text_stream = read_text_file_generator(args.file)

    # 2. Stream Chunks & Process
    chunk_stream = stream_chunks(text_stream)

    print(f"⚡ Processing stream...")

    chunk_count = 0
    success_count = 0

    # We iterate over chunks as they are generated
    for i, chunk in enumerate(chunk_stream):
        chunk_count += 1
        payload = {
            "content": chunk,
            "tags": ["ingested", f"{ext_clean}_import", base_name, f"chunk_{i}"]
        }

        # 3. Output or Send Immediately
        if args.url and not args.dry_run:
            try:
                response = httpx.post(args.url, json=payload, timeout=10.0)
                if response.status_code == 200:
                    success_count += 1
                else:
                    print(f"  - Chunk {i} failed ❌ ({response.status_code})")
            except Exception as e:
                print(f"  - Chunk {i} Error: {e}")
        else:
            # Dry run: just print summary every 10 chunks or full JSON for first few
            if i < 3:
                print(json.dumps(payload, indent=2, ensure_ascii=False))
            elif i == 3:
                print("... (streaming continues) ...")

    print(f"\n🏁 Completed. Processed {chunk_count} chunks.")
    if args.url and not args.dry_run:
        print(f"🚀 Successfully sent: {success_count}/{chunk_count}")

if __name__ == "__main__":
    main()
